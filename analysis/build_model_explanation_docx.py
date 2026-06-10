from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"D:/AAA Learning/AAA本科毕业设计/材料")
OUT_PATH = OUT_DIR / "时空网格态势推演与复合瓶颈指数模型说明.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=11, bold=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_format(paragraph, first_line=True, before=0, after=6, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=14 if level == 1 else 12, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_math(doc, text):
    # Word uses Office Math Markup Language (OMML). The equation text is placed
    # inside a math paragraph so that it opens as a native Word equation object.
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    xml = f"""
    <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
                 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <m:oMath>
        <m:r>
          <m:rPr><m:sty m:val="p"/></m:rPr>
          <m:t>{escaped}</m:t>
        </m:r>
      </m:oMath>
    </m:oMathPara>
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(parse_xml(xml))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=10, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("基于时空网格的客流态势推演与复合瓶颈指数识别模型")
    set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("面向高铁枢纽突发客流行人仿真的方法说明")
    set_run_font(r, east_asia="楷体", ascii_font="Times New Roman", size=11)

    add_heading(doc, "1 基于时空网格的客流态势推演模型", 1)

    add_heading(doc, "1.1 研究场景中的问题指向", 2)
    add_body_paragraph(
        doc,
        "高铁枢纽出站通道中的突发客流具有显著的时变性、方向性和局部汇聚特征。以 Burst 工况为例，"
        "仿真总客流量为 4699 人，4 个出站口分别承担 1104、1295、1437 和 863 人。"
        "该类客流并非在空间上均匀展开，而是由多个出站口在短时间内集中汇入东西向走廊，"
        "容易形成局部高占用、局部滞留以及拥堵波沿走廊传播等现象。"
        "如果仅以个体轨迹或单帧散点图描述仿真结果，虽然可以呈现行人位置，却难以回答三个关键问题："
        "客流压力在何处累积、拥堵过程在何时形成、局部压力是否随时间向上下游传导。"
    )
    add_body_paragraph(
        doc,
        "因此，本文提出基于时空网格的客流态势推演模型，其目的并不是替代 MARL 行人仿真模型，"
        "而是将仿真产生的大量微观状态转化为可解释、可比较、可定位的宏观态势场。"
        "该模型面向高铁枢纽突发客流场景，重点解决微观仿真结果难以直接支撑运行诊断的问题，"
        "为后续瓶颈识别提供连续的时空状态基础。"
    )

    add_heading(doc, "1.2 时空网格状态表达", 2)
    add_body_paragraph(
        doc,
        "在模型实现中，出站通道被离散为二维空间网格。东西向走廊方向记为 X，横向通道方向记为 Z。"
        "本研究数据中，X 向共有 30 个网格单元，编号为 0 至 29；Z 向共有 3 个网格单元，编号为 0 至 2。"
        "对任一时刻 t，网格单元可表示为："
    )
    add_math(doc, "G_{i,j},  i=0,1,...,29;  j=0,1,2")
    add_body_paragraph(
        doc,
        "每个网格单元记录人数、密度、平均速度以及速度分量，形成局部客流状态向量："
    )
    add_math(doc, "S_{i,j}(t)=[n_{i,j}(t), ρ_{i,j}(t), v_{i,j}(t), v^x_{i,j}(t), v^z_{i,j}(t)]")
    add_body_paragraph(
        doc,
        "其中，n 表示网格内行人数，ρ 表示局部密度，v 表示平均速度，"
        "v^x 与 v^z 分别表示东西向和横向速度分量。"
        "与仅记录个体坐标相比，该表达能够把局部空间占用、运动强度和运动方向统一到同一时空框架中。"
    )

    add_heading(doc, "1.3 沿走廊方向的态势聚合", 2)
    add_body_paragraph(
        doc,
        "由于本文关注的是东西向出站通道的运行状态，代码中进一步将 Z 向网格进行聚合，"
        "得到沿 X 方向的一维走廊态势。该处理保留了瓶颈在走廊方向上的空间定位能力，"
        "同时降低了横向细节对整体趋势判断的干扰。聚合后的人数、密度与加权速度分别为："
    )
    add_math(doc, "N_i(t)=Σ_j n_{i,j}(t)")
    add_math(doc, "ρ_i(t)=Σ_j ρ_{i,j}(t)")
    add_math(doc, "v̄_i(t)=Σ_j n_{i,j}(t)v_{i,j}(t) / Σ_j n_{i,j}(t)")
    add_body_paragraph(
        doc,
        "其中，N_i(t) 表示第 i 个 X 向位置的总人数，ρ_i(t) 表示该位置的聚合密度，"
        "v̄_i(t) 表示按人数加权后的平均速度。当某一 X 向位置没有行人时，速度不参与加权计算。"
        "代码中的 groupby(['time_bin','cell_x']) 操作即对应上述聚合过程。"
    )
    add_body_paragraph(
        doc,
        "为评估通道整体承载压力，模型进一步计算实时活跃人数："
    )
    add_math(doc, "A(t)=Σ_i N_i(t)")
    add_body_paragraph(
        doc,
        "在 Burst 工况下，A(t) 的峰值为 412 人，出现在 333.59 s。"
        "这一指标表明突发客流并非简单地表现为总人数增加，而是在特定时间段内形成了显著的通道占用峰值。"
        "态势推演模型由此能够刻画客流从进入、累积、峰值到消散的全过程。"
    )

    add_heading(doc, "1.4 数据甄别与仿真状态修正", 2)
    add_body_paragraph(
        doc,
        "高铁枢纽突发客流仿真中，局部网格可能出现人数很少但速度异常偏高的情况。"
        "若不加甄别，这类稀疏网格速度会干扰速度损失和瓶颈指数计算。"
        "因此，模型在态势聚合之前设置速度有效性判别，将超过 3.0 的速度值视为异常速度，"
        "在后续加权速度计算中剔除："
    )
    add_math(doc, "v^*_{i,j}(t)=v_{i,j}(t),  if 0<v_{i,j}(t)<=3.0;  otherwise null")
    add_body_paragraph(
        doc,
        "以 Burst 数据为例，原始数据结构完整，共 151560 行，1684 个时间步，"
        "不存在缺失值、重复 time-cell 记录、负人数或负密度；但 mean_speed 大于 3.0 的非空网格共有 1849 行，"
        "最大速度达到 12.4203。该异常主要出现在单人稀疏网格中，说明在突发客流仿真中，"
        "状态推演模型不仅需要汇总信息，还需要承担数据甄别与状态清洗功能。"
    )
    add_table(
        doc,
        ["统计项", "Burst 工况结果", "模型处理含义"],
        [
            ["总行数", "151560", "网格时序数据规模较大，需要转化为态势场"],
            ["时间步数", "1684", "支持连续过程推演"],
            ["活跃人数峰值", "412（333.59 s）", "识别通道承载压力峰值"],
            ["最大密度", "0.2813", "作为局部高占用状态的重要参照"],
            ["速度异常行数", "1849（>3.0）", "在速度加权和瓶颈识别前剔除"],
        ],
        widths=[3.6, 4.0, 8.0],
    )

    add_heading(doc, "1.5 模型作用小结", 2)
    add_body_paragraph(
        doc,
        "综上，基于时空网格的客流态势推演模型针对的是突发客流仿真结果“数据量大、局部性强、过程难解释”的问题。"
        "其核心作用包括三点：第一，将微观行人运动状态转化为 X 向连续态势场；第二，识别通道压力峰值及其时空分布；"
        "第三，为后续瓶颈识别提供密度、人数、速度和增长趋势等可计算变量。"
        "该模型使高铁枢纽突发客流仿真从“轨迹展示”转向“运行状态诊断”。"
    )

    add_heading(doc, "2 基于复合瓶颈指数的运行瓶颈识别模型", 1)

    add_heading(doc, "2.1 研究场景中的瓶颈识别难点", 2)
    add_body_paragraph(
        doc,
        "在高铁枢纽突发客流场景中，运行瓶颈并不等同于某一时刻密度最高的网格。"
        "一方面，局部密度升高可能只是瞬时汇聚，不一定造成持续运行障碍；另一方面，"
        "某些位置即使密度不是全局最高，也可能因速度下降或密度持续增长而形成排队传播。"
        "因此，若仅依据单一密度阈值识别瓶颈，容易出现两类偏差：将短时波动误判为瓶颈，"
        "或忽略持续时间较长但速度损失不剧烈的运行压力区。"
    )
    add_body_paragraph(
        doc,
        "基于此，本文构建复合瓶颈指数，将密度压力、速度损失和密度增长趋势纳入统一评价框架。"
        "该模型的目标是从时空态势场中识别真正影响出站通道运行效率的区域，并区分“严格瓶颈事件”"
        "与“持续拥堵热点”。"
    )

    add_heading(doc, "2.2 复合瓶颈指数构建", 2)
    add_body_paragraph(
        doc,
        "首先，定义密度风险项。考虑到不同客流工况下密度水平差异较大，模型没有采用固定密度阈值，"
        "而是基于当前工况自动计算密度阈值："
    )
    add_math(doc, "ρ_0=max(P_90(ρ_i(t)), 0.6ρ_max)")
    add_body_paragraph(
        doc,
        "其中，P_90(ρ_i(t)) 表示非空网格聚合密度的 90 分位数，ρ_max 表示最大聚合密度。"
        "Burst 工况下，模型得到 ρ_0=0.18。密度风险项定义为："
    )
    add_math(doc, "R_{ρ,i}(t)=clip(ρ_i(t)/ρ_0, 0, 1)")
    add_body_paragraph(
        doc,
        "其次，定义速度损失项。代码中设定自由行走速度 v_f=1.3，当局部平均速度低于自由速度时，"
        "该位置被认为存在运行阻滞："
    )
    add_math(doc, "R_{v,i}(t)=clip(1-v̄_i(t)/v_f, 0, 1),  v_f=1.3")
    add_body_paragraph(
        doc,
        "再次，定义密度增长项，用于描述拥堵是否处于形成或加剧过程："
    )
    add_math(doc, "R_{g,i}(t)=clip((ρ_i(t)-ρ_i(t-Δt))/ρ_0, 0, 1)")
    add_body_paragraph(
        doc,
        "在此基础上，复合瓶颈指数定义为："
    )
    add_math(doc, "B_i(t)=0.50R_{ρ,i}(t)+0.35R_{v,i}(t)+0.15R_{g,i}(t)")
    add_body_paragraph(
        doc,
        "权重设置体现了模型对高铁枢纽出站通道运行机理的理解：密度压力是瓶颈形成的主要表现，"
        "速度损失反映局部通行受阻程度，密度增长则用于捕捉瓶颈形成过程中的动态变化。"
        "代码中设置瓶颈阈值 θ_b=0.62，并要求连续持续 L=3 个时间步，以过滤单步随机扰动："
    )
    add_math(doc, "B_i(t)≥θ_b,  θ_b=0.62,  persistence≥L,  L=3")

    add_heading(doc, "2.3 严格瓶颈与候选拥堵热点的区分", 2)
    add_body_paragraph(
        doc,
        "突发客流场景中的运行风险具有层次性。严格瓶颈强调“高风险且持续”，更适合识别局部停滞或明显通行受阻事件；"
        "候选拥堵热点则强调“高占用且持续”，更适合描述突发客流下长期承压的通道区段。"
        "为避免将短时速度为零的个别网格等同于主瓶颈，模型进一步引入候选热点指数："
    )
    add_math(doc, "H_i(t)=0.65R_{ρ,i}(t)+0.35R_{N,i}(t)")
    add_math(doc, "R_{N,i}(t)=clip(N_i(t)/N_0, 0, 1)")
    add_body_paragraph(
        doc,
        "其中，N_0 为网格人数的高分位参照值。该指标更强调密度与人数占用的共同作用，"
        "适用于识别突发客流条件下持续存在的运行压力区。"
    )

    add_heading(doc, "2.4 Burst 工况下的瓶颈识别结果", 2)
    add_body_paragraph(
        doc,
        "基于上述模型，Burst 工况下严格瓶颈事件共识别出 2 个，均位于 cell_x=5，"
        "发生在 174.19 s 和 174.69 s，局部速度为 0，瓶颈指数分别约为 0.703 和 0.714。"
        "然而，该事件持续时间短，空间范围有限，因此更适合作为短时停滞点，而不宜作为主运行瓶颈。"
    )
    add_body_paragraph(
        doc,
        "候选拥堵热点结果则揭示了更加稳定的运行瓶颈结构。Burst 工况下，主要瓶颈集中在 cell_x=6、7、8。"
        "其中，cell_x=8 平均人数为 40.97，最大人数为 51，平均密度为 0.2305，最大密度为 0.2869；"
        "cell_x=7 平均人数为 38.39，最大人数为 53，平均密度为 0.2159，最大密度为 0.2981；"
        "cell_x=6 平均人数为 33.62，最大人数为 46，平均密度为 0.1892，最大密度为 0.2588。"
    )
    add_table(
        doc,
        ["位置", "热点持续步数", "时间范围", "平均人数", "最大人数", "平均密度", "最大密度", "解释"],
        [
            ["cell_x=8", "1146", "39.47-619.10 s", "40.97", "51", "0.2305", "0.2869", "密度核心"],
            ["cell_x=7", "1159", "46.98-634.18 s", "38.39", "53", "0.2159", "0.2981", "持续性最强"],
            ["cell_x=6", "1109", "65.28-641.69 s", "33.62", "46", "0.1892", "0.2588", "上游堆积区"],
            ["cell_x=9", "13", "102.10-582.80 s", "25.77", "28", "0.1449", "0.1575", "下游扩散区"],
            ["cell_x=19", "21", "159.59-569.20 s", "25.57", "28", "0.1439", "0.1575", "次级压力点"],
        ],
        widths=[2.2, 2.3, 3.1, 2.0, 2.0, 2.0, 2.0, 3.0],
    )
    add_body_paragraph(
        doc,
        "上述结果表明，Burst 工况下的主瓶颈不是一个孤立网格，而是由 cell_x=6-8 构成的持续性高压带。"
        "其中 cell_x=8 体现最高密度，cell_x=7 体现最长持续性，cell_x=6 则表现为上游排队堆积。"
        "从时间维度看，核心瓶颈从约 40-65 s 开始形成，并持续至 620-640 s 左右，"
        "说明突发客流下通道压力具有显著的长时段维持特征。"
    )

    add_heading(doc, "2.5 模型作用小结", 2)
    add_body_paragraph(
        doc,
        "复合瓶颈指数模型针对的是突发客流仿真中“瓶颈判定标准单一、短时波动与持续瓶颈难以区分”的问题。"
        "通过将密度、速度损失和密度增长纳入同一指数体系，模型能够识别局部运行受阻事件；"
        "通过候选热点指数，模型又能够揭示高客流条件下持续承压的空间区段。"
        "结合 Burst 工况结果可以认为，cell_x=6-8 是出站通道的主运行瓶颈区，"
        "cell_x=9-10 为下游扩散区，cell_x=19 附近存在次级压力点。"
        "这一结论不仅给出了瓶颈的位置，也说明了瓶颈的形成时段、持续性与空间传导关系，"
        "能够为高铁枢纽突发客流下的闸机开放、通道组织和客流诱导提供定量依据。"
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
