from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


SRC = Path(r"D:/AAA Learning/AAA本科毕业设计/材料/毕业论文.docx")
OUT = Path(r"D:/AAA Learning/AAA本科毕业设计/材料/毕业论文_第四章改写_STG_CBI.docx")


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def find_para(doc: Document, needle: str, start: int = 0) -> int:
    for i, para in enumerate(doc.paragraphs[start:], start):
        if needle in para.text:
            return i
    raise ValueError(f"Cannot find paragraph containing {needle!r}")


def find_heading(doc: Document, needle: str, start: int = 0, level: str | None = None) -> int:
    for i, para in enumerate(doc.paragraphs[start:], start):
        style = para.style.name if para.style is not None else ""
        if needle in para.text and style.startswith("Heading"):
            if level is None or style == level:
                return i
    raise ValueError(f"Cannot find heading containing {needle!r}")


def insert_before(anchor, text: str, style: str | None = None):
    para = anchor.insert_paragraph_before(text)
    if style:
        para.style = style
    if style == "Normal":
        para.paragraph_format.first_line_indent = Pt(21)
    return para


def chapter4_content():
    return [
        ("第四章 高铁枢纽行人流态势推演与运行瓶颈识别模型", "Heading 1"),
        ("4.1 方法框架与数据适配性", "Heading 2"),
        ("本章围绕出站通道东西向走廊内的客流运行状态，构建“STG-CBI：时空网格复合瓶颈指数法”（Spatio-Temporal Grid Composite Bottleneck Index, STG-CBI）。该方法不是直接对单个行人轨迹进行展示，而是将 Unity 与 MARL 仿真输出的实时位置和速度信息转化为可解释的网格化状态，再进一步形成时空态势矩阵和运行瓶颈判别结果。", "Normal"),
        ("STG-CBI 的计算链条包括六个环节：第一，基于出站通道空间坐标进行网格化状态感知；第二，按采样时间构建密度、人数、速度和速度分量的时空状态矩阵；第三，沿走廊 X 方向聚合 Z 向网格，得到东西向通道的主导运行态势；第四，计算密度风险、速度损失和密度增长风险，形成复合瓶颈指数；第五，引入连续采样步约束，区分严格运行瓶颈和相对拥堵热点；第六，通过热力图、快照图、单元剖面图和事件表对瓶颈位置、持续时间和扩散范围进行解释。", "Normal"),
        ("本章方法直接对应 Python 分析脚本 crowd_bottleneck_visualization.py 的实际实现。脚本输入字段为 scenario、time、cell_x、cell_z、count、density、mean_speed、mean_vx 和 mean_vz。Normal 工况共有 64080 条网格状态记录，Burst 工况共有 151560 条网格状态记录，二者均满足完整网格结构，无缺失值、无重复 time-cell 记录、无负人数和负密度。", "Normal"),
        ("在数据甄别过程中发现，两类数据均存在少量稀疏网格的异常高速值。Normal 工况 mean_speed 最高为 12.2064，超过 3.0 的非空网格有 241 行；Burst 工况 mean_speed 最高为 12.4203，超过 3.0 的非空网格有 1849 行。考虑到这些异常多出现在单人稀疏单元，本文将 3.0 作为速度质检上限，超过该上限的速度记录不参与加权平均速度和瓶颈速度损失计算，但仍导出为 quality_high_speed_rows.csv 供数据追溯。", "Normal"),
        ("4.2 网格化客流状态感知模型", "Heading 2"),
        ("设出站通道被划分为 X 向 30 个单元、Z 向 3 个单元。任一采样时刻 t 的网格单元记为 c=(x,z)，对应状态向量为 s_c(t)={n_c(t), rho_c(t), v_c(t), v_x,c(t), v_z,c(t)}。其中 n_c(t) 为网格人数，rho_c(t) 为网格密度，v_c(t) 为平均速度，v_x,c(t) 和 v_z,c(t) 分别为速度在走廊长度方向和宽度方向上的分量。", "Normal"),
        ("由于论文研究对象是东西向出站走廊的运行瓶颈，分析阶段将 Z 向 3 个网格聚合到同一 X 单元上。聚合后 X 单元的人数为 N_x(t)=sum_z n_(x,z)(t)，聚合密度为 R_x(t)=sum_z rho_(x,z)(t)。平均速度采用人数加权方式计算：V_x(t)=sum_z n_(x,z)(t)v_(x,z)(t)/sum_z n_(x,z)(t)。当某一 Z 向网格人数为 0 或速度超过质检上限时，该速度不参与加权平均。", "Normal"),
        ("该处理方式使模型既保留走廊宽度方向的信息，又能形成沿 X 方向连续可读的瓶颈演化图。对于本文数据而言，X=0 至 X=29 覆盖完整出站通道，Z=0 至 Z=2 表示通道宽度方向的三条统计带，因此可以支持通道纵向热力图和典型时刻二维快照两类可视化。", "Normal"),
        ("4.3 基于时空网格矩阵的客流态势推演模型", "Heading 2"),
        ("客流态势推演采用时空网格矩阵表达。对每一类指标 m，可构建矩阵 M_m=[m_x(t)]，其中矩阵行表示采样时间，矩阵列表示走廊 X 单元。本文实际输出的核心矩阵包括人数矩阵 C(t,x)、密度矩阵 D(t,x)、速度矩阵 V(t,x) 和瓶颈指数矩阵 B(t,x)。", "Normal"),
        ("在推演意义上，模型关注的是相邻采样时刻之间的状态变化，而不是重新训练一个预测模型。密度增长项 Delta D(t,x)=D(t,x)-D(t-1,x) 用于判断某一位置是否处于持续积压过程；速度变化和平均速度水平用于判断通行能力是否下降。这样可以将 MARL 微观仿真输出转化为“哪里人变多、哪里速度变慢、哪里风险持续增强”的宏观态势。", "Normal"),
        ("态势推演的可视化包括四类图：一是 active_count_curve，用于观察场景内活跃人数随时间变化；二是 density_heatmap_x_time，用于识别高密度区在空间和时间上的扩展；三是 speed_heatmap_x_time，用于识别低速或停滞区；四是 corridor_situation_snapshots，用网格人数底图叠加速度方向箭头，展示典型时刻的二维空间态势。", "Normal"),
        ("4.4 基于复合瓶颈指数的运行瓶颈识别模型", "Heading 2"),
        ("本文将运行瓶颈定义为在一定持续时间内同时表现出较高局部密度、明显速度损失和密度增长趋势的空间单元。为避免仅依靠单一密度阈值造成误判，STG-CBI 构建三个归一化风险项。", "Normal"),
        ("密度风险定义为 R_d(t,x)=clip(D(t,x)/theta_d,0,1)。其中 theta_d 为密度阈值。由于 Normal 和 Burst 两类工况密度量级不同，本文采用自适应密度阈值 theta_d=max(P90(D_nonzero),0.6*max(D_nonzero),1e-6)。实际计算中，Normal 工况 theta_d=0.1351，Burst 工况 theta_d=0.1800。", "Normal"),
        ("速度损失定义为 R_v(t,x)=clip(1-V(t,x)/v_f,0,1)，其中 v_f=1.30 表示自由行走速度。若该 X 单元当前人数为 0，则速度损失记为 0，避免空网格被误识别为低速区域。", "Normal"),
        ("密度增长风险定义为 R_g(t,x)=clip((D(t,x)-D(t-1,x))/theta_d,0,1)。当密度下降或保持稳定时，增长风险为 0；当密度在相邻采样周期内明显增加时，增长风险升高。", "Normal"),
        ("综合瓶颈指数定义为 B(t,x)=0.50R_d(t,x)+0.35R_v(t,x)+0.15R_g(t,x)。其中密度风险权重最高，用于体现高占用对通道容量的压力；速度损失次之，用于体现通行能力下降；密度增长风险用于捕捉瓶颈形成过程中的短时积压趋势。", "Normal"),
        ("为减少瞬时噪声带来的误判，本文设置瓶颈指数阈值 beta=0.62，并要求连续 persistence_steps=3 个采样周期满足 B(t,x)>=beta。由于采样间隔约为 0.5 s，严格运行瓶颈至少需要持续约 1.5 s 才会被输出为 bottleneck_events.csv。", "Normal"),
        ("4.5 候选拥堵热点识别与瓶颈解释", "Heading 2"),
        ("严格瓶颈判定强调“高密度+低速度+持续性”。但在 Normal 或有序 Burst 场景中，局部区域可能长期高占用却未出现明显速度损失。为避免遗漏这类潜在瓶颈，本文进一步设置候选拥堵热点指标 H(t,x)=0.65R_d(t,x)+0.35R_c(t,x)。其中 R_c(t,x)=clip(C(t,x)/P95(C),0,1)，表示相对人数压力。若 H(t,x)>=0.75 且连续 3 个采样周期成立，则输出为 candidate_hotspot_events.csv。", "Normal"),
        ("因此，本文将结果分为两层解释：严格运行瓶颈表示已经发生的低速拥堵状态；候选拥堵热点表示持续高占用、可能诱发瓶颈的风险区。该双层判别更契合高铁站出站通道场景，因为正常组织下乘客可能保持移动，但局部高压区仍然是后续调控和设施优化的重点。", "Normal"),
        ("在 Normal 工况下，严格运行瓶颈事件为 0，说明 548 人、3 出口条件下未形成严重持续低速瓶颈；但候选热点有 1469 个时空单元，主要集中在 X=6 至 X=10，其中 X=8 为最高风险核心。在 Burst 工况下，活跃人数峰值达到 412，候选热点增加到 3494 个时空单元，主热点集中在 X=6 至 X=8，其中 X=8 为密度核心，X=7 持续性最强，X=6 表现为上游排队堆积区。严格瓶颈事件仅在 X=5 附近出现 2 个短时事件，主要由局部速度为 0 触发，解释时不宜将其作为主瓶颈，而应将 X=6 至 X=8 作为 Burst 工况下的持续运行瓶颈带。", "Normal"),
        ("4.6 可视化效果、解释性与数据需求检查", "Heading 2"),
        ("从当前数据结构看，STG-CBI 的可解释性较好。密度热力图能够展示高占用区的形成和扩散，速度热力图能够揭示低速或停滞位置，瓶颈指数热力图能够把密度、速度和增长趋势合成为统一风险面，典型时刻快照能够直观呈现二维空间上的人数分布和速度方向。对于 Burst 工况，X=6 至 X=8 的高压带在热力图和快照图中均清晰可见，符合运行瓶颈识别目标。", "Normal"),
        ("当前可视化也存在边界。CSV 文件是网格聚合态势数据，不包含行人唯一编号、完整个体轨迹、到达闸机时间、实际排队等待时间和闸机逐口服务率。因此，本文可以可靠识别“哪里拥挤、何时拥挤、是否持续、是否伴随速度损失”，但不能直接绘制每个乘客的完整轨迹线，也不能精确计算单个乘客排队时长、OD 路径选择比例或每个闸机的真实服务效率。", "Normal"),
        ("若后续需要进一步增强可视化和论文说服力，建议补充采集以下数据：第一，agent_id、time、x、z、speed、target_gate 级别的个体轨迹表，用于绘制轨迹束、路径选择和滞留时间；第二，每名乘客 spawn_time、gate_arrival_time、exit_id、gate_id，用于计算通行时间和排队等待时间；第三，闸机逐口通过日志和服务时间，用于评估东西闸机负载均衡；第四，走廊几何尺寸、障碍物位置和闸机物理坐标，用于将 cell_x 映射为真实空间位置；第五，真实监测或人工统计客流，用于对仿真密度和速度进行校准。", "Normal"),
        ("4.7 本章小结", "Heading 2"),
        ("本章提出了面向高铁站出站通道的 STG-CBI 时空网格复合瓶颈指数法。该方法以 Unity/MARL 仿真输出的实时网格状态为基础，通过 X-Z 网格化状态感知、Z 向聚合、时空矩阵构建、复合瓶颈指数计算和连续性判别，实现了从微观仿真数据到宏观客流态势推演和运行瓶颈识别的转换。", "Normal"),
        ("与仅依靠密度阈值的方法相比，STG-CBI 同时考虑密度风险、速度损失和密度增长趋势，并额外设置候选拥堵热点层，用于解释高占用但尚未明显停滞的潜在瓶颈区域。该方法能够较好地区分 Normal 工况下的相对热点和 Burst 工况下的持续瓶颈带，具有较强的可解释性和可视化表达能力。", "Normal"),
        ("结合实际运行结果，Normal 工况未形成严重持续瓶颈，但 X=8 附近存在相对拥堵热点；Burst 工况下 X=6 至 X=8 形成持续高压带，是出站通道运行瓶颈识别和后续组织优化的重点区域。后续章节将在该模型基础上进一步分析不同客流场景下的仿真结果和瓶颈演化特征。", "Normal"),
    ]


def main():
    doc = Document(SRC)

    ch4_start = find_heading(doc, "第四章", level="Heading 1")
    ch5_start = find_heading(doc, "第五章", ch4_start + 1, level="Heading 1")
    ch5_anchor = doc.paragraphs[ch5_start]
    old_ch4_paras = list(doc.paragraphs[ch4_start:ch5_start])

    misplaced_46_paras = []
    try:
        old_46 = find_heading(doc, "4.6 仿真", ch5_start)
        next_51 = find_heading(doc, "5.1", old_46 + 1)
        misplaced_46_paras = list(doc.paragraphs[old_46:next_51])
    except ValueError:
        pass

    for para in old_ch4_paras:
        delete_paragraph(para)
    for para in misplaced_46_paras:
        delete_paragraph(para)

    for text, style in chapter4_content():
        insert_before(ch5_anchor, text, style)

    for section in doc.sections:
        section.top_margin = section.top_margin

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
